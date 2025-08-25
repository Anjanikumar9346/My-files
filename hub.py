import os
import requests
import base64
from flask import Flask, request, jsonify
from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document

load_dotenv("hub.env")

app = Flask(__name__)

GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_REPO_OWNER = os.getenv("GITHUB_REPO_OWNER")
GITHUB_REPO_NAME = os.getenv("GITHUB_REPO_NAME")

HEADERS = {
    "Authorization": f"token {GITHUB_TOKEN}"
}

VALID_EXTENSIONS = [".pdf", ".txt", ".docx"]

# ---------------------------
# GitHub API Helpers
# ---------------------------

def list_all_branches():
    url = f"https://api.github.com/repos/{GITHUB_REPO_OWNER}/{GITHUB_REPO_NAME}/branches"
    response = requests.get(url, headers=HEADERS)
    if response.status_code != 200:
        return []
    return [branch["name"] for branch in response.json()]

def fetch_files_from_branch(branch_name):
    url = f"https://api.github.com/repos/{GITHUB_REPO_OWNER}/{GITHUB_REPO_NAME}/git/trees/{branch_name}?recursive=1"
    response = requests.get(url, headers=HEADERS)
    if response.status_code != 200:
        return []
    tree = response.json().get("tree", [])
    return [
        item["path"]
        for item in tree
        if item["type"] == "blob"
        and os.path.splitext(item["path"])[1].lower() in VALID_EXTENSIONS
        and os.path.basename(item["path"]).lower() != "readme.md"
    ]

def fetch_all_valid_files():
    branches = list_all_branches()
    all_files = set()
    for branch in branches:
        files = fetch_files_from_branch(branch)
        all_files.update(files)
    return list(all_files)

def get_file_content_from_github(file_path):
    url = f"https://api.github.com/repos/{GITHUB_REPO_OWNER}/{GITHUB_REPO_NAME}/contents/{file_path}"
    response = requests.get(url, headers=HEADERS)
    if response.status_code != 200:
        return None, response.status_code
    content = base64.b64decode(response.json()["content"])
    return content, 200

# ---------------------------
# Content Extraction Helpers
# ---------------------------

def extract_description(content, ext):
    try:
        if ext == ".txt":
            return content.decode().split("\n")[0]
        elif ext == ".pdf":
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            return text.strip().split("\n")[0]
        elif ext == ".docx":
            with open("temp.docx", "wb") as f:
                f.write(content)
            doc = Document("temp.docx")
            return doc.paragraphs[0].text if doc.paragraphs else ""
        else:
            return "Unsupported file"
    except Exception as e:
        return f"Error reading file: {str(e)}"

# ---------------------------
# API 1: Get all files with descriptions
# ---------------------------

@app.route("/get-all-files", methods=["GET"])
def get_all_files():
    files = fetch_all_valid_files()
    output = []

    for file_path in files:
        file_name = os.path.basename(file_path)
        ext = os.path.splitext(file_name)[1].lower()
        content, status = get_file_content_from_github(file_path)
        if content:
            desc = extract_description(content, ext)
        else:
            desc = "Unable to fetch"
        output.append({
            "file_name": file_name,
            "description": desc
        })

    return jsonify(output)

# ---------------------------
# API 2: Get file content by filename
# ---------------------------

@app.route("/get-file-content", methods=["POST"])
def get_file_content():
    data = request.json
    filename = data.get("file_name")
    if not filename:
        return jsonify({"error": "file_name is required"}), 400

    files = fetch_all_valid_files()
    match = next((f for f in files if os.path.basename(f) == filename), None)

    if not match:
        return jsonify({"error": "File not found"}), 404

    ext = os.path.splitext(filename)[1].lower()
    content, status = get_file_content_from_github(match)
    if status != 200 or content is None:
        return jsonify({"error": "Could not fetch file"}), 500

    if ext == ".txt":
        return jsonify({
            "file_name": filename,
            "content": content.decode()
        })
    elif ext == ".pdf":
        doc = fitz.open(stream=content, filetype="pdf")
        text = "".join([page.get_text() for page in doc])
        return jsonify({
            "file_name": filename,
            "content": text.strip()
        })
    elif ext == ".docx":
        with open("temp.docx", "wb") as f:
            f.write(content)
        doc = Document("temp.docx")
        text = "\n".join(p.text for p in doc.paragraphs)
        return jsonify({
            "file_name": filename,
            "content": text.strip()
        })
    else:
        return jsonify({"error": "Unsupported file type"}), 400

# ---------------------------
# App Runner
# ---------------------------

if __name__ == "__main__":
    app.run(debug=True)